import hashlib
import io
import unittest
import zipfile

from server.knowledge_parsing import build_document_ir, document_ir_to_markdown, persisted_block_rows


def build_ir(filename, raw, extracted_text, mime_type="text/plain", document_id="doc-1"):
    return build_document_ir(
        document_id=document_id,
        filename=filename,
        raw=raw,
        extracted_text=extracted_text,
        source_mime=mime_type,
        source_sha256=hashlib.sha256(raw).hexdigest(),
        parser_version="fixture-v1",
    )


def minimal_xlsx() -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr(
            "xl/workbook.xml",
            """<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
            <sheets><sheet name="预算" sheetId="1" r:id="rId1"/></sheets></workbook>""",
        )
        archive.writestr(
            "xl/_rels/workbook.xml.rels",
            """<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
            <Relationship Id="rId1" Target="worksheets/sheet1.xml"/></Relationships>""",
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            """<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
            <sheetData><row r="1"><c r="A1" t="inlineStr"><is><t>项目</t></is></c>
            <c r="B1" t="inlineStr"><is><t>金额</t></is></c></row>
            <row r="2"><c r="A2" t="inlineStr"><is><t>研发</t></is></c>
            <c r="B2"><v>120</v></c></row></sheetData></worksheet>""",
        )
    return buffer.getvalue()


class KnowledgeParsingTests(unittest.TestCase):
    def test_markdown_preserves_structure_and_is_stable(self):
        text = """# 产品方案

第一段说明。

## 范围

- Web
- API

| 项目 | 状态 |
| --- | --- |
| 解析 | 完成 |
"""
        first = build_ir("plan.md", text.encode(), text, "text/markdown")
        second = build_ir("plan.md", text.encode(), text, "text/markdown")
        self.assertEqual(
            [block.block_id for block in first.blocks],
            [block.block_id for block in second.blocks],
        )
        self.assertEqual(
            [block.block_type for block in first.blocks],
            ["heading", "paragraph", "heading", "list", "table"],
        )
        self.assertEqual(first.blocks[-1].section_path, ("产品方案", "范围"))
        markdown = document_ir_to_markdown(first)
        self.assertIn("# 产品方案", markdown)
        self.assertIn("- Web", markdown)
        self.assertIn("| 项目 | 状态 |", markdown)

    def test_pdf_pages_are_traceable(self):
        extracted = "【PDF 第 1 页】\n封面\n\n【PDF 第 2 页】\n正文"
        document_ir = build_ir("guide.pdf", b"%PDF-fixture", extracted, "application/pdf")
        self.assertEqual([block.source_location["page"] for block in document_ir.blocks], [1, 2])
        self.assertEqual(document_ir.blocks[1].section_path, ("第 2 页",))

    def test_xlsx_preserves_sheet_and_cell_range(self):
        raw = minimal_xlsx()
        document_ir = build_ir(
            "budget.xlsx",
            raw,
            "【工作表：预算】\n项目 | 金额\n研发 | 120",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        self.assertEqual(len(document_ir.blocks), 1)
        block = document_ir.blocks[0]
        self.assertEqual(block.block_type, "sheet")
        self.assertEqual(block.source_location, {"sheet": "预算", "cell_range": "A1:B2"})
        self.assertIn("研发 | 120", block.text)

    def test_image_ocr_records_source(self):
        document_ir = build_ir("scan.png", b"png-fixture", "识别后的文字", "image/png")
        self.assertEqual(document_ir.blocks[0].block_type, "image_ocr")
        self.assertEqual(document_ir.blocks[0].source_location["image"], "scan.png")

    def test_docx_preserves_heading_list_and_table(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx unavailable")
        document = Document()
        document.add_heading("项目说明", level=1)
        document.add_paragraph("目标内容")
        document.add_paragraph("第一项", style="List Bullet")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "字段"
        table.rows[0].cells[1].text = "值"
        buffer = io.BytesIO()
        document.save(buffer)
        raw = buffer.getvalue()
        document_ir = build_ir(
            "guide.docx",
            raw,
            "项目说明\n目标内容\n第一项\n字段\n值",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertEqual(
            [block.block_type for block in document_ir.blocks],
            ["heading", "paragraph", "list", "table"],
        )
        self.assertEqual(document_ir.blocks[-1].source_location["table"], 1)

    def test_persisted_rows_are_content_addressed(self):
        document_ir = build_ir("note.txt", b"alpha", "alpha")
        rows = persisted_block_rows(document_ir, "run-1", 100)
        self.assertEqual(rows[0]["document_id"], "doc-1")
        self.assertEqual(rows[0]["char_count"], 5)
        self.assertEqual(rows[0]["content_sha256"], hashlib.sha256(b"alpha").hexdigest())


if __name__ == "__main__":
    unittest.main()
